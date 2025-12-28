"""
File Parser Service for NodeRAG API
====================================

Handles parsing of binary files (PDF, DOCX, TXT) to extract text content.
"""

import io
from typing import Optional
from pathlib import Path


class FileParser:
    """Service for parsing various file formats to extract text"""
    
    @staticmethod
    def parse_pdf(file_bytes: bytes) -> str:
        """
        Extract text from PDF using PyPDF2
        
        Args:
            file_bytes: PDF file content as bytes
            
        Returns:
            Extracted text content
        """
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF parsing. Install with: pip install PyPDF2")
        
        try:
            pdf_file = io.BytesIO(file_bytes)
            reader = PdfReader(pdf_file)
            
            text_content = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)
            
            return "\n\n".join(text_content)
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    @staticmethod
    def parse_docx(file_bytes: bytes) -> str:
        """
        Extract text from DOCX using python-docx
        
        Args:
            file_bytes: DOCX file content as bytes
            
        Returns:
            Extracted text content
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
        
        try:
            docx_file = io.BytesIO(file_bytes)
            document = Document(docx_file)
            
            text_content = []
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_content.append(row_text)
            
            return "\n\n".join(text_content)
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    @staticmethod
    def parse_txt(file_bytes: bytes) -> str:
        """
        Decode text file
        
        Args:
            file_bytes: Text file content as bytes
            
        Returns:
            Decoded text content
        """
        try:
            # Try UTF-8 first
            return file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1 (which accepts all byte sequences)
            try:
                return file_bytes.decode('latin-1')
            except Exception as e:
                raise ValueError(f"Failed to decode text file: {str(e)}")
    
    @staticmethod
    def detect_file_type(filename: str, file_bytes: bytes) -> str:
        """
        Detect file type from filename and content
        
        Args:
            filename: Original filename
            file_bytes: File content as bytes
            
        Returns:
            File type: 'pdf', 'docx', or 'txt'
        """
        # First, try by extension
        extension = Path(filename).suffix.lower()
        
        if extension == '.pdf':
            return 'pdf'
        elif extension in ['.docx', '.doc']:
            return 'docx'
        elif extension in ['.txt', '.md', '.text']:
            return 'txt'
        
        # Try magic bytes detection
        if file_bytes.startswith(b'%PDF'):
            return 'pdf'
        elif file_bytes.startswith(b'PK\x03\x04'):  # ZIP signature (DOCX is a ZIP file)
            # Check if it's a DOCX by looking for document.xml
            return 'docx'
        
        # Default to txt
        return 'txt'
    
    @classmethod
    def parse_file(cls, file_bytes: bytes, filename: str) -> str:
        """
        Main entry point - detect file type and parse accordingly
        
        Args:
            file_bytes: File content as bytes
            filename: Original filename
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file type is unsupported or parsing fails
        """
        file_type = cls.detect_file_type(filename, file_bytes)
        
        if file_type == 'pdf':
            return cls.parse_pdf(file_bytes)
        elif file_type == 'docx':
            return cls.parse_docx(file_bytes)
        elif file_type == 'txt':
            return cls.parse_txt(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")


# Convenience function for direct use
def parse_file(file_bytes: bytes, filename: str) -> str:
    """
    Parse a file and extract its text content
    
    Args:
        file_bytes: File content as bytes
        filename: Original filename
        
    Returns:
        Extracted text content
    """
    return FileParser.parse_file(file_bytes, filename)

